from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "BazaarHelper_User_Guide.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(32, 45, 58)
MUTED = RGBColor(91, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF4D6"


def set_font(run, size: float, *, bold: bool = False, color: RGBColor = DARK) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        set_font(lead, 11, bold=True)
        body = paragraph.add_run(text[len(bold_prefix) :])
        set_font(body, 11)
    else:
        run = paragraph.add_run(text)
        set_font(run, 11)


def add_step(doc: Document, title: str, detail: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    style_paragraph(paragraph, after=5)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    title_run = paragraph.add_run(title + " ")
    set_font(title_run, 11, bold=True, color=BLUE)
    detail_run = paragraph.add_run(detail)
    set_font(detail_run, 11)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    style_paragraph(paragraph, after=4)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    set_font(paragraph.add_run(text), 11)


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0)
    set_font(paragraph.add_run(label + "："), 11, bold=True, color=BLUE)
    set_font(paragraph.add_run(text), 11)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10),
        "Heading 2": (13, 14, 7),
        "Heading 3": (12, 10, 5),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Number", "List Bullet"):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(paragraph.add_run("BazaarHelper 内部测试使用教程"), 9, color=MUTED)


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(5)
    set_font(title.add_run("BazaarHelper 使用教程"), 28, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run("The Bazaar 对局事件分析助手 | 内部测试版"), 13, color=MUTED)

    add_callout(
        doc,
        "最快开始",
        "完整解压文件夹 → 运行 install_plugin.bat → 启动游戏 → 运行 start.bat。",
        LIGHT_GOLD,
    )

    add_heading(doc, "一、收到文件后", 1)
    add_bullet(doc, "完整解压 BazaarHelper 文件夹，不要只复制 BazaarHelper.exe。")
    add_bullet(doc, "不要删除 _internal、data、examples 或 bepinex_plugin 文件夹。")
    add_bullet(doc, "建议解压到桌面或其他普通文件夹，程序文件夹之后可以自由移动。")

    add_heading(doc, "二、首次安装", 1)
    add_step(doc, "关闭游戏。", "安装或更新插件前，先完全退出 The Bazaar。")
    add_step(doc, "运行安装脚本。", "双击 install_plugin.bat。")
    add_step(
        doc,
        "输入游戏目录。",
        r"输入包含 BepInEx 文件夹的 The Bazaar 安装目录，例如 D:\Steam\steamapps\common\The Bazaar。",
    )
    add_step(doc, "确认安装完成。", "窗口显示安装完成后按任意键关闭。")
    add_step(doc, "启动游戏。", "进入一局游戏，插件会自动读取当前状态。")

    add_callout(
        doc,
        "安装提示",
        "如果脚本提示没有检测到 BepInEx，需要先为游戏安装 BepInEx，再重新运行安装脚本。",
    )

    add_heading(doc, "三、日常使用", 1)
    add_step(doc, "先启动游戏。", "进入当前对局。")
    add_step(doc, "启动助手。", "双击 start.bat，浏览器会自动打开 http://127.0.0.1:8765。")
    add_step(doc, "进入事件页面。", "助手会读取当前事件并给出推荐。")
    add_step(doc, "等待自动刷新。", "页面每 3 秒自动读取一次状态，也可以点击“刷新”。")
    add_step(doc, "需要时点击 AI 分析。", "AI 分析不会阻塞普通刷新。")

    add_heading(doc, "四、AI 功能", 1)
    add_body(doc, "内部测试包已附带测试 API Key，首次运行 start.bat 时会自动配置，无需手动填写。")
    add_body(doc, "如需更换 Key，可双击 set_ai_key.bat，在记事本中粘贴新的 Key，保存后重新启动助手。")
    add_callout(doc, "内部测试", "此版本及其中的 API Key 仅供指定测试人员使用。", LIGHT_GOLD)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "五、阶段定义", 1)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    widths = [Inches(1.8), Inches(4.7)]
    headers = ("阶段", "天数")
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        set_cell_margins(cell)
        shade_cell(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(paragraph.add_run(text), 11, bold=True, color=BLUE)

    for stage, days in (("前期", "第 1-5 天"), ("中期", "第 6-9 天"), ("后期", "第 10 天及以后")):
        cells = table.add_row().cells
        for index, text in enumerate((stage, days)):
            cells[index].width = widths[index]
            set_cell_margins(cells[index])
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(paragraph.add_run(text), 11, bold=index == 0)

    add_heading(doc, "六、移动文件夹或更换电脑", 1)
    add_bullet(doc, "在同一台电脑上移动或改名 BazaarHelper 文件夹，不需要重新安装插件。")
    add_bullet(doc, "更换电脑后，需要在新电脑重新运行 install_plugin.bat。")
    add_bullet(doc, "游戏安装位置改变后，需要重新运行 install_plugin.bat。")
    add_bullet(doc, r"运行数据保存在 %LOCALAPPDATA%\BazaarHelper\runtime，不依赖助手文件夹位置。")

    add_heading(doc, "七、不刷新时的排查顺序", 1)
    add_step(doc, "确认游戏正在运行。", "助手只能读取当前游戏状态。")
    add_step(doc, "确认进入了事件页面。", "停留在主菜单或加载界面时可能没有可分析事件。")
    add_step(doc, "关闭旧助手再重开。", "再次双击 start.bat 会自动结束旧进程并启动最新版。")
    add_step(doc, "重新安装插件。", "退出游戏后再次运行 install_plugin.bat，然后重启游戏。")
    add_step(doc, "检查完整文件夹。", "确认 data、_internal 和 bepinex_plugin 没有缺失。")
    add_step(doc, "保留现场信息。", "记录游戏天数、当前事件，并截取助手页面，方便定位。")

    add_heading(doc, "八、常见问题", 1)
    add_body(doc, "浏览器没有自动打开：手动访问 http://127.0.0.1:8765。", bold_prefix="浏览器没有自动打开：")
    add_body(doc, "AI 分析失败：重新运行 set_ai_key.bat 检查 Key，然后重启助手。", bold_prefix="AI 分析失败：")
    add_body(doc, "移动目录后能否继续使用：可以，运行数据使用固定用户目录。", bold_prefix="移动目录后能否继续使用：")
    add_body(doc, "更新版本：用新文件夹替换旧版本，退出游戏后重新运行 install_plugin.bat。", bold_prefix="更新版本：")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
